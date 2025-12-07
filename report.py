from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from datetime import datetime
from llm_client import call_llm, get_system_prompt
import io


def build_report_context(current_summary: Dict, previous_summary: Dict, 
                        inspection_meta: Dict, delta: Dict = None) -> Dict:
    """
    Собирает контекст для генерации отчёта
    """
    context = {
        'customer': 'Заказчик (не указан)',
        'pipeline_name': inspection_meta.get('pipeline_name', 'N/A'),
        'segment_km': inspection_meta.get('segment_km', 'N/A'),
        'diameter_mm': inspection_meta.get('diameter_mm', 'N/A'),
        'method': inspection_meta.get('method', 'N/A'),
        'inspection_date': str(inspection_meta.get('start_date', datetime.now().strftime('%Y-%m-%d'))),
        'total_defects': current_summary['overview']['total_defects'],
        'high_risk_count': current_summary['by_risk']['High'],
        'medium_risk_count': current_summary['by_risk']['Medium'],
        'low_risk_count': current_summary['by_risk']['Low'],
        'defect_types': current_summary['by_type'],
        'statistics': current_summary['statistics'],
        'has_previous_inspection': delta and delta.get('has_previous', False),
        'changes': delta if delta else {}
    }
    
    return context


def generate_report_texts(context: Dict) -> Dict:
    """
    Генерирует текстовые разделы отчёта через LLM
    """
    system_prompt = get_system_prompt()
    
    # 1. Краткое заключение
    user_prompt_summary = """Сформулируй краткое заключение (2-3 абзаца) по результатам обследования трубопровода.
Включи:
- Общее количество обнаруженных дефектов
- Распределение по классам риска
- Основные выводы о состоянии трубопровода"""
    
    summary_text = call_llm(system_prompt, user_prompt_summary, context)
    
    # 2. Описание результатов обследования
    user_prompt_results = """Опиши результаты обследования (3-4 абзаца).
Включи:
- Метод обследования и охваченный участок
- Детализацию по типам обнаруженных аномалий
- Статистику по глубине дефектов и ERF
- Анализ критичных дефектов высокого риска"""
    
    results_text = call_llm(system_prompt, user_prompt_results, context)
    
    # 3. Сравнение с предыдущей инспекцией (если есть)
    comparison_text = ""
    if context.get('has_previous_inspection'):
        user_prompt_comparison = """Опиши изменения относительно предыдущей инспекции (2-3 абзаца).
Включи:
- Изменение общего количества дефектов
- Динамику по классам риска
- Общую тенденцию (улучшение/ухудшение состояния)"""
        
        comparison_text = call_llm(system_prompt, user_prompt_comparison, context)
    else:
        comparison_text = "Данные предыдущей инспекции отсутствуют."
    
    # 4. Рекомендации
    user_prompt_recommendations = """Сформулируй рекомендации по ремонту и дальнейшему мониторингу (3-4 пункта).
Включи:
- Приоритетные ремонтные работы для дефектов высокого риска
- Рекомендации по наблюдению за дефектами среднего риска
- Предложения по периодичности следующих обследований"""
    
    recommendations_text = call_llm(system_prompt, user_prompt_recommendations, context)
    
    return {
        'summary': summary_text,
        'results': results_text,
        'comparison': comparison_text,
        'recommendations': recommendations_text
    }


def render_docx(context: Dict, texts: Dict) -> bytes:
    """
    Создаёт Word-документ с отчётом
    """
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок
    title = doc.add_heading('Заключительный отчёт об обследовании трубопровода', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"{context['pipeline_name']}, участок {context['segment_km']} км")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Таблица с метаданными
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    metadata = [
        ('Заказчик:', context.get('customer', 'N/A')),
        ('Трубопровод:', context['pipeline_name']),
        ('Участок:', f"{context['segment_km']} км"),
        ('Диаметр:', f"{context['diameter_mm']} мм"),
        ('Метод обследования:', context['method']),
        ('ID отчёта:', f"REP-{datetime.now().strftime('%Y%m%d-%H%M%S')}"),
        ('Дата:', context['inspection_date'])
    ]
    
    for i, (label, value) in enumerate(metadata):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 1. Краткое заключение
    doc.add_heading('1. Краткое заключение', level=1)
    doc.add_paragraph(texts['summary'])
    
    # 2. Результаты обследования
    doc.add_heading('2. Результаты обследования', level=1)
    doc.add_paragraph(texts['results'])
    
    # Статистика в виде списка
    doc.add_paragraph('Сводная статистика:', style='Heading 2')
    stats_list = [
        f"Всего обнаружено дефектов: {context['total_defects']}",
        f"Высокий риск: {context['high_risk_count']}",
        f"Средний риск: {context['medium_risk_count']}",
        f"Низкий риск: {context['low_risk_count']}"
    ]
    
    for stat in stats_list:
        p = doc.add_paragraph(stat, style='List Bullet')
    
    # 3. Сравнение с предыдущей инспекцией
    doc.add_heading('3. Динамика изменений', level=1)
    doc.add_paragraph(texts['comparison'])
    
    # 4. Рекомендации
    doc.add_heading('4. Рекомендации', level=1)
    doc.add_paragraph(texts['recommendations'])
    
    # Сохраняем в байты
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()

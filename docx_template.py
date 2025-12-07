"""
Модуль для заполнения template.docx заранее подготовленным шаблоном
"""

from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import os
from typing import Dict


def fill_template_docx(template_path: str, 
                       context: Dict, 
                       texts: Dict,
                       scheme_image_path: str,
                       output_path: str) -> str:
    """
    Заполняет заранее подготовленный template.docx
    
    Args:
        template_path: путь к template.docx (ОБЯЗАТЕЛЕН!)
        context: контекст с данными
        texts: сгенерированные тексты от LLM
        scheme_image_path: путь к scheme_with_defects.png
        output_path: куда сохранить результат
    
    Returns:
        путь к созданному файлу
    """
    
    # Проверяем наличие шаблона
    if not os.path.exists(template_path):
        raise FileNotFoundError(
            f"❌ Файл {template_path} не найден!\n"
            f"Пожалуйста, поместите файл template.docx в корень проекта.\n"
            f"Это должен быть шаблон Word с логотипами и печатями."
        )
    
    # Загружаем шаблон
    doc = Document(template_path)
    
    # Словарь для замены плейсхолдеров
    replacements = {
        '{{ЗАКАЗЧИК}}': context.get('customer', '________'),
        '{{ТРУБОПРОВОД}}': context.get('pipeline_name', 'N/A'),
        '{{УЧАСТОК}}': context.get('segment_km', 'N/A'),
        '{{ПРОТЯЖЕННОСТЬ}}': f"{context.get('length_km', '1.0')} км",
        '{{ДИАМЕТР}}': f"{context.get('diameter_mm', 'N/A')} мм",
        '{{ID_ОТЧЕТА}}': f"REP-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
        '{{ДАТА}}': str(context.get('inspection_date', datetime.now().strftime('%Y-%m-%d'))),
        
        # Статистика
        '{{ВСЕГО_ДЕФЕКТОВ}}': str(context.get('total_defects', 0)),
        '{{ВЫСОКИЙ_РИСК}}': str(context.get('high_risk_count', 0)),
        '{{СРЕДНИЙ_РИСК}}': str(context.get('medium_risk_count', 0)),
        '{{НИЗКИЙ_РИСК}}': str(context.get('low_risk_count', 0)),
        
        # Тексты от LLM
        '{{ЗАКЛЮЧЕНИЕ}}': texts.get('summary', ''),
        '{{РЕЗУЛЬТАТЫ}}': texts.get('results', ''),
        '{{СРАВНЕНИЕ}}': texts.get('comparison', ''),
        '{{РЕКОМЕНДАЦИИ}}': texts.get('recommendations', ''),
    }
    
    # Заменяем плейсхолдеры в параграфах
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Заменяем плейсхолдеры в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    
    # Ищем место для вставки схемы (маркер {{СХЕМА}})
    scheme_inserted = False
    for i, paragraph in enumerate(doc.paragraphs):
        if '{{СХЕМА}}' in paragraph.text:
            # Удаляем маркер
            paragraph.text = ''
            
            # Вставляем изображение
            if scheme_image_path and os.path.exists(scheme_image_path):
                run = paragraph.add_run()
                run.add_picture(scheme_image_path, width=Inches(6.5))
                
                # Центрируем
                paragraph.alignment = 1  # CENTER
                scheme_inserted = True
            break
    
    # Если маркер не найден, добавляем схему в конец
    if not scheme_inserted and scheme_image_path and os.path.exists(scheme_image_path):
        doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_picture(scheme_image_path, width=Inches(6.5))
        paragraph.alignment = 1
    
    # Сохраняем
    doc.save(output_path)
    return output_path


def create_blank_template(output_path: str = "template.docx"):
    """
    Создаёт пустой шаблон с плейсхолдерами для примера
    
    Используйте эту функцию один раз, чтобы создать базовый шаблон,
    затем отредактируйте его вручную, добавив логотипы и печати
    """
    doc = Document()
    
    # Шапка
    doc.add_paragraph("CENTER PROFESSIONAL ENGINEERING")
    doc.add_paragraph("100012, Модау1, Казахстан")
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()
    
    # УТВЕРЖДАЮ
    approve = doc.add_paragraph("УТВЕРЖДАЮ:")
    approve.alignment = 2  # RIGHT
    approve2 = doc.add_paragraph("Директор ТОО «Center Professional Engineering»")
    approve2.alignment = 2
    approve3 = doc.add_paragraph("_" * 25)
    approve3.alignment = 2
    approve4 = doc.add_paragraph("дата")
    approve4.alignment = 2
    doc.add_paragraph()
    
    # Заголовок
    title = doc.add_heading("Заключительный отчёт", level=0)
    title.alignment = 1  # CENTER
    subtitle = doc.add_paragraph("об обследовании трубопровода")
    subtitle.alignment = 1
    doc.add_paragraph()
    
    # Таблица
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    fields = [
        ('Заказчик', '{{ЗАКАЗЧИК}}'),
        ('Трубопровод', '{{ТРУБОПРОВОД}}'),
        ('Участок', '{{УЧАСТОК}}'),
        ('Протяжённость', '{{ПРОТЯЖЕННОСТЬ}}'),
        ('Диаметр', '{{ДИАМЕТР}}'),
        ('ID отчёта', '{{ID_ОТЧЕТА}}'),
        ('Дата', '{{ДАТА}}')
    ]
    
    for i, (label, placeholder) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = placeholder
    
    doc.add_page_break()
    
    # Разделы
    doc.add_heading("1. Краткое заключение", level=1)
    doc.add_paragraph("{{ЗАКЛЮЧЕНИЕ}}")
    doc.add_paragraph()
    
    doc.add_heading("2. Результаты обследования", level=1)
    doc.add_paragraph("{{РЕЗУЛЬТАТЫ}}")
    doc.add_paragraph()
    
    doc.add_heading("2.1 Сводная статистика", level=2)
    doc.add_paragraph(f"Всего дефектов: {{{{ВСЕГО_ДЕФЕКТОВ}}}}")
    doc.add_paragraph(f"Высокий риск: {{{{ВЫСОКИЙ_РИСК}}}}")
    doc.add_paragraph(f"Средний риск: {{{{СРЕДНИЙ_РИСК}}}}")
    doc.add_paragraph(f"Низкий риск: {{{{НИЗКИЙ_РИСК}}}}")
    doc.add_paragraph()
    
    doc.add_heading("2.2 Схема расположения дефектов", level=2)
    doc.add_paragraph("{{СХЕМА}}")
    doc.add_paragraph()
    
    doc.add_page_break()
    
    doc.add_heading("3. Динамика изменений", level=1)
    doc.add_paragraph("{{СРАВНЕНИЕ}}")
    doc.add_paragraph()
    
    doc.add_heading("4. Рекомендации", level=1)
    doc.add_paragraph("{{РЕКОМЕНДАЦИИ}}")
    
    doc.save(output_path)
    print(f"✅ Создан базовый шаблон: {output_path}")
    print("Отредактируйте его вручную, добавив логотипы и печати!")
    
    return output_path
